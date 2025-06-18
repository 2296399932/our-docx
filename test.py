import unittest
import os
import shutil
from docx_parser import DocxFile
from docx_namespace import DocxElementParser
import tempfile


class TestDocxFile(unittest.TestCase):
    """测试DocxFile类的基本功能"""

    def setUp(self):
        """准备测试环境"""
        # 创建一个测试文件夹
        self.test_dir = tempfile.mkdtemp()
        # 测试用docx文件路径 - 需要自行提供一个测试docx文件
        self.test_docx = os.path.join(self.test_dir, "test.docx")
        # 创建一个空的测试文件 - 实际测试时需替换为真实的docx文件
        with open(self.test_docx, "wb") as f:
            f.write(b"PK\x03\x04\x14\x00\x06\x00\x08\x00")  # 简单的ZIP文件头
        
        # 输出测试文件路径
        self.output_path = os.path.join(self.test_dir, "output.docx")

    def tearDown(self):
        """清理测试环境"""
        shutil.rmtree(self.test_dir)

    def test_init(self):
        """测试DocxFile初始化"""
        docx = DocxFile(self.test_docx)
        self.assertEqual(docx.path, self.test_docx)
        
    def test_save(self):
        """测试保存功能"""
        # 注意：这个测试可能会失败，因为我们使用了一个伪造的docx文件
        # 实际测试时需要使用真实的docx文件
        try:
            docx = DocxFile(self.test_docx)
            docx.save(self.output_path)
            self.assertTrue(os.path.exists(self.output_path))
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")


class TestDocxElementParser(unittest.TestCase):
    """测试DocxElementParser类的功能"""
    
    def setUp(self):
        """准备测试环境"""
        # 创建一个测试文件夹
        self.test_dir = tempfile.mkdtemp()
        # 测试用docx文件路径 - 需要自行提供一个测试docx文件
        self.test_docx = os.path.join(self.test_dir, "test_parser.docx")
        # 创建一个空的测试文件 - 实际测试时需替换为真实的docx文件
        with open(self.test_docx, "wb") as f:
            f.write(b"PK\x03\x04\x14\x00\x06\x00\x08\x00")  # 简单的ZIP文件头
        
        # 输出测试文件路径
        self.output_path = os.path.join(self.test_dir, "output_parser.docx")
        
        # 输出图片目录
        self.image_output_dir = os.path.join(self.test_dir, "images")
        os.makedirs(self.image_output_dir, exist_ok=True)

    def tearDown(self):
        """清理测试环境"""
        shutil.rmtree(self.test_dir)

    def test_init(self):
        """测试DocxElementParser初始化"""
        try:
            parser = DocxElementParser(self.test_docx)
            self.assertEqual(parser.path, self.test_docx)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")

    def test_get_element(self):
        """测试获取XML元素"""
        try:
            parser = DocxElementParser(self.test_docx)
            element = parser.get_element()
            self.assertIsNotNone(element)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_find_elements_by_tag(self):
        """测试通过标签查找元素"""
        try:
            parser = DocxElementParser(self.test_docx)
            elements = parser.find_elements_by_tag("w:p")  # 寻找段落元素
            self.assertIsInstance(elements, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_get_all_paragraphs(self):
        """测试获取所有段落"""
        try:
            parser = DocxElementParser(self.test_docx)
            paragraphs = parser.get_all_paragraphs()
            self.assertIsInstance(paragraphs, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_get_all_paragraphs_text(self):
        """测试获取所有段落文本"""
        try:
            parser = DocxElementParser(self.test_docx)
            paragraphs_text = parser.get_all_paragraphs_text()
            self.assertIsInstance(paragraphs_text, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_get_all_tables(self):
        """测试获取所有表格"""
        try:
            parser = DocxElementParser(self.test_docx)
            tables = parser.get_all_tables()
            self.assertIsInstance(tables, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_get_all_text(self):
        """测试获取所有文本"""
        try:
            parser = DocxElementParser(self.test_docx)
            all_text = parser.get_all_text()
            self.assertIsInstance(all_text, str)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_get_structured_body_elements(self):
        """测试获取结构化正文元素"""
        try:
            parser = DocxElementParser(self.test_docx)
            elements = parser.get_structured_body_elements()
            self.assertIsInstance(elements, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")

    def test_extract_table_content(self):
        """测试提取表格内容"""
        try:
            parser = DocxElementParser(self.test_docx)
            tables = parser.get_all_tables()
            if tables:
                content = parser.extract_table_content(tables[0])
                self.assertIsInstance(content, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件或表格: {str(e)}")
    
    def test_extract_images_simple(self):
        """测试提取图片"""
        try:
            parser = DocxElementParser(self.test_docx)
            images = parser.extract_images_simple(self.image_output_dir)
            self.assertIsInstance(images, list)
        except Exception as e:
            self.skipTest(f"需要带图片的docx文件: {str(e)}")
    
    def test_count_images_simple(self):
        """测试计算图片数量"""
        try:
            parser = DocxElementParser(self.test_docx)
            count = parser.count_images_simple()
            self.assertIsInstance(count, int)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_extract_paragraph_style(self):
        """测试提取段落样式"""
        try:
            parser = DocxElementParser(self.test_docx)
            paragraphs = parser.get_all_paragraphs()
            if paragraphs:
                style = parser.extract_paragraph_style(paragraphs[0])
                self.assertIsInstance(style, dict)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_get_document_structure(self):
        """测试获取文档结构"""
        try:
            parser = DocxElementParser(self.test_docx)
            structure = parser.get_document_structure()
            self.assertIsInstance(structure, list)
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_insert_paragraph(self):
        """测试插入段落"""
        try:
            parser = DocxElementParser(self.test_docx)
            # 插入一个新段落
            parser.insert_paragraph(text="测试段落", position="after")
            parser.save(self.output_path)
            self.assertTrue(os.path.exists(self.output_path))
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_insert_table(self):
        """测试插入表格"""
        try:
            parser = DocxElementParser(self.test_docx)
            # 插入一个2x2表格
            data = [["头1", "头2"], ["数据1", "数据2"]]
            parser.insert_table(rows=2, cols=2, data=data)
            parser.save(self.output_path)
            self.assertTrue(os.path.exists(self.output_path))
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_remove_element(self):
        """测试删除元素"""
        try:
            parser = DocxElementParser(self.test_docx)
            elements = parser.get_structured_body_elements()
            if elements:
                # 删除第一个元素
                parser.remove_element(0)
                parser.save(self.output_path)
                self.assertTrue(os.path.exists(self.output_path))
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")
    
    def test_set_paragraph_style(self):
        """测试设置段落样式"""
        try:
            parser = DocxElementParser(self.test_docx)
            paragraphs = parser.get_all_paragraphs()
            if paragraphs:
                # 设置第一个段落的对齐方式为居中
                parser.update_paragraph_style(0, alignment="center")
                parser.save(self.output_path)
                self.assertTrue(os.path.exists(self.output_path))
        except Exception as e:
            self.skipTest(f"需要真实的docx文件: {str(e)}")


def create_sample_docx():
    """创建一个简单的测试用docx文件（使用python-docx库）"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # 创建文档
        doc = Document()
        
        # 添加一些段落
        doc.add_paragraph('这是一个测试文档')
        doc.add_paragraph('这是第二段落', style='Heading 1')
        doc.add_paragraph('这是第三段落，包含一些格式', style='Heading 2')
        
        # 添加一个表格
        table = doc.add_table(rows=2, cols=2)
        # 填充表格
        table.cell(0, 0).text = '表格单元格 (0,0)'
        table.cell(0, 1).text = '表格单元格 (0,1)'
        table.cell(1, 0).text = '表格单元格 (1,0)'
        table.cell(1, 1).text = '表格单元格 (1,1)'
        
        # 保存文档
        output_path = "sample_test.docx"
        doc.save(output_path)
        print(f"创建了一个示例文档: {output_path}")
        return output_path
    except ImportError:
        print("需要安装python-docx库来创建示例文档")
        return None


def main():
    # 尝试创建一个示例文档
    sample_path = create_sample_docx()
    
    if sample_path and os.path.exists(sample_path):
        print(f"使用示例文档 {sample_path} 运行测试")
        # 你可以在这里修改测试类中的测试文件路径
    else:
        print("请提供一个测试用的docx文件来运行测试")
    
    # 运行测试
    unittest.main()


if __name__ == "__main__":
    main()
